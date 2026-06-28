import io
import os
import re
import numpy as np
import torch
from docx import Document
from transformers import pipeline, AutoTokenizer, AutoModelForSequenceClassification
from duckduckgo_search import DDGS

# --- AI DETECTOR CLASS WRAPPER ---

class AIDetectorEngine:
    """
    Manages loading and running the OpenAI RoBERTa detector model.
    """
    def __init__(self, model_name="roberta-base-openai-detector"):
        self.device = 0 if torch.cuda.is_available() else -1
        self.model_name = model_name
        
        # Load Hugging Face pipeline
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
        self.model = AutoModelForSequenceClassification.from_pretrained(self.model_name)
        self.detector = pipeline(
            "text-classification", 
            model=self.model, 
            tokenizer=self.tokenizer, 
            device=self.device
        )


# --- FORMAT CHECKING & EXTRACTION UTILITIES ---

def extract_text_from_docx(file_source) -> str:
    """
    Extracts plain text sentences from either a path string 
    or an in-memory BytesIO object.
    """
    try:
        doc = Document(file_source)
        full_text = []
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if len(stripped) >= 10:
                full_text.append(stripped)
        return " ".join(full_text)
    except Exception:
        return ""


def extract_sentences(text: str) -> list:
    """Splits a single block of text into valid evaluation sentences."""
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)
    valid_sentences = []
    for s in sentences:
        s_clean = s.strip()
        # Only evaluate sentences containing 8 or more words for search matching
        if len(s_clean.split()) >= 8:
            valid_sentences.append(s_clean)
    return valid_sentences


# --- WEB CRAWLER / SEARCH INTEGRATION ---

def is_sentence_on_web(sentence: str) -> tuple:
    """
    Queries DuckDuckGo using double quotes for exact phrase searches.
    Returns (True, URL) if an exact phrase match is found, else (False, None).
    """
    query = f'"{sentence}"'
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=1))
            if len(results) > 0:
                return True, results[0].get('href', 'Unknown Source')
    except Exception:
        pass
    return False, None


# --- STREAM-COMPATIBLE INTEGRITY RUNNER ---

def run_document_integrity_check(input_file, detector_engine=None, progress_callback=None) -> dict:
    """
    Runs the comprehensive integrity assessment (AI classification and Web plagiarism check).
    
    Args:
        input_file: Path string or BytesIO instance containing the Word Document [1].
        detector_engine: An optional pre-initialized instance of AIDetectorEngine.
        progress_callback: Callable function(step_name, current_progress_percent).
        
    Returns:
        dictionary containing structured results, scores, matched sources, and logs.
    """
    # Initialize detector if not provided to allow for caching
    if detector_engine is None:
        detector_engine = AIDetectorEngine()

    doc = Document(input_file)
    target_text = extract_text_from_docx(input_file)
    
    # Initialize results structures
    results = {
        "success": False,
        "paragraphs_scanned": 0,
        "sentences_checked": 0,
        "ai_score": 0.0,
        "ai_verdict": "N/A",
        "similarity_index": 0.0,
        "similarity_verdict": "N/A",
        "matched_sources": [],
        "matched_sentences": [],
        "logs": []
    }

    if not target_text:
        results["logs"].append("ERROR: Document text is unreadable or empty.")
        return results

    # --- Part A: AI Detection ---
    results["logs"].append("Step 1/2: Processing AI content classification...")
    total_ai_score = 0
    paragraph_count = 0
    total_paragraphs = len(doc.paragraphs)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if len(text) < 50:
            continue

        if progress_callback:
            progress_callback("AI Detection", int((i / total_paragraphs) * 100))

        # Truncate to standard token boundaries
        text_chunk = text[:1500]

        try:
            prediction = detector_engine.detector(text_chunk)[0]
            label = prediction['label']
            score = prediction['score']

            if label == 'Real':
                human_prob = score * 100
            else:
                human_prob = (1 - score) * 100

            total_ai_score += human_prob
            paragraph_count += 1
        except Exception as e:
            results["logs"].append(f"AI scanning issue on paragraph index {i}: {str(e)}")

    if paragraph_count > 0:
        avg_human_score = total_ai_score / paragraph_count
        ai_score = 100 - avg_human_score
    else:
        ai_score = 0.0

    # Determine AI Verdict
    if ai_score < 15:
        ai_verdict = "Likely Human-Written (Safe)"
    elif ai_score < 45:
        ai_verdict = "Mixed/Paraphrased (Moderate Risk)"
    else:
        ai_verdict = "High AI Content (High Risk)"

    results["paragraphs_scanned"] = paragraph_count
    results["ai_score"] = round(ai_score, 2)
    results["ai_verdict"] = ai_verdict

    # --- Part B: Plagiarism Web Scan ---
    results["logs"].append("Step 2/2: Sampling and scanning sentences online...")
    sentences = extract_sentences(target_text)
    total_sentences = len(sentences)

    # Intelligently sample up to 15 sentences to help avoid search engine rate limits
    if total_sentences > 15:
        indices = np.linspace(0, total_sentences - 1, 15, dtype=int)
        sampled_sentences = [sentences[idx] for idx in indices]
    else:
        sampled_sentences = sentences

    matched_count = 0
    sources_found = set()
    matched_sentences_list = []
    total_sampled = len(sampled_sentences)

    for s_idx, sentence in enumerate(sampled_sentences):
        if progress_callback:
            progress_callback("Web Scan", int((s_idx / total_sampled) * 100))

        found, url = is_sentence_on_web(sentence)
        if found:
            matched_count += 1
            sources_found.add(url)
            matched_sentences_list.append({"sentence": sentence, "url": url})
            results["logs"].append(f"Plagiarism match: \"{sentence[:45]}...\" -> {url}")

    # Similarity Index Calculation
    if len(sampled_sentences) > 0:
        similarity_index = (matched_count / len(sampled_sentences)) * 100
    else:
        similarity_index = 0.0

    # Determine Similarity Verdict
    if similarity_index < 15:
        sim_verdict = "Low Similarity (Safe)"
    elif similarity_index < 30:
        sim_verdict = "Moderate Similarity (Acceptable, check citations)"
    else:
        sim_verdict = "High Similarity (High Risk of Plagiarism flag)"

    results["sentences_checked"] = len(sampled_sentences)
    results["similarity_index"] = round(similarity_index, 2)
    results["similarity_verdict"] = sim_verdict
    results["matched_sources"] = list(sources_found)
    results["matched_sentences"] = matched_sentences_list
    results["success"] = True

    return results
