import io
import re
from docx import Document
from docx.text.paragraph import Paragraph

# Headings indicating the bibliography/references section
REF_HEADINGS = [
    "references", "bibliography", "works cited", "literature cited",
    "sources", "reference list", "literature"
]

def remove_references_and_after(doc) -> tuple:
    """
    Finds the references/bibliography section heading and removes it
    along with all elements (paragraphs, tables, etc.) that follow it.
    Returns (bool, str) indicating success status and log message.
    """
    body = doc._body._element
    ref_element = None

    # Iterate through child elements of the document body
    for child in body:
        if child.tag.endswith('p'):  # It's a paragraph element
            p = Paragraph(child, doc)
            text = p.text.strip().lower()
            # Clean punctuation from heading for matching (e.g., "References." -> "references")
            clean_text = re.sub(r'[^\w\s]', '', text).strip()
            if clean_text in REF_HEADINGS:
                ref_element = child
                break

    if ref_element is not None:
        children = list(body)
        ref_idx = children.index(ref_element)
        heading_text = Paragraph(ref_element, doc).text
        log_msg = f"Found References section starting with heading: '{heading_text}'"
        
        # Delete from the end of the body down to the references heading
        for idx in range(len(children) - 1, ref_idx - 1, -1):
            body.remove(children[idx])
        return True, log_msg + " - References section and subsequent content successfully removed."
    
    return False, "References/Bibliography section header not detected. No trailing content was removed."


def replace_text_and_preserve_style(p, new_text):
    """
    Overwrites the paragraph text while retaining paragraph properties
    (like margins, line spacing, alignments, and heading level)
    and applying the basic font style of the original paragraph.
    """
    style = p.style

    if not p.runs:
        p.text = new_text
        p.style = style
        return

    # Capture styling of the first run to apply to the updated text
    first_run = p.runs[0]
    font_name = first_run.font.name
    font_size = first_run.font.size
    bold = first_run.bold
    italic = first_run.italic
    color = first_run.font.color.rgb if first_run.font.color else None

    p.text = ""  # Clear existing text

    new_run = p.add_run(new_text)
    new_run.font.name = font_name
    new_run.font.size = font_size
    if bold is not None:
        new_run.bold = bold
    if italic is not None:
        new_run.italic = italic
    if color:
        new_run.font.color.rgb = color

    p.style = style


def clean_paragraph(p) -> bool:
    """
    Identifies and removes parenthetical, narrative, and bracketed citations,
    as well as numerical superscript citations.
    Returns True if paragraph was modified, False otherwise.
    """
    modified = False
    
    # 1. Handle superscript citations (Vancouver numeric superscript)
    for run in list(p.runs):
        if run.font.superscript:
            text = run.text.strip()
            # If the superscript run contains only numbers, commas, dashes, or spaces
            if re.match(r'^[\d\s,\-–—]+$', text):
                run.text = ""
                modified = True

    text = p.text
    if not text.strip():
        return modified

    original_text = text

    # Pattern A: Numbered bracket citations, e.g., [1], [1, 2], [1-3]
    numbered_pattern = r'\[\s*\d+(?:\s*[\s\-,]\s*\d+)*\s*\]'
    text = re.sub(numbered_pattern, '', text)

    # Pattern B: Narrative citations with year inside parentheses, e.g., (2020) or (2020, p. 15)
    narrative_pattern = r'\(\s*\d{4}[a-z]?(?:\s*,\s*(?:p|pp)\.?\s*\d+(?:-\d+)?)?\s*\)'
    text = re.sub(narrative_pattern, '', text)

    # Pattern C: Parenthetical author-date citations, e.g., (Smith, 2020), (Smith & Jones, 2018), or multi-citations (Smith, 2020; Doe, 2019)
    author_date_pattern = r'\(\s*[A-Za-z\s&,\.\-–′’\']+\d{4}[a-z]?(?:\s*,\s*(?:p|pp)\.?\s*\d+(?:-\d+)?)?\s*(?:;\s*[A-Za-z\s&,\.\-–′’\']+\d{4}[a-z]?(?:\s*,\s*(?:p|pp)\.?\s*\d+(?:-\d+)?)?\s*)*\)'
    text = re.sub(author_date_pattern, '', text)

    # 2. Cleanup spacing artifacts left behind by deletions
    # Remove spacing before punctuation (e.g. "word (Citation)." becomes "word . " -> "word.")
    text = re.sub(r'\s+(?=[.,;:?!])', '', text)
    # Replace multiple spaces with a single space
    text = re.sub(r' {2,}', ' ', text)
    text = text.strip()

    # 3. Update paragraph if matches were found
    if text != original_text:
        replace_text_and_preserve_style(p, text)
        return True
        
    return modified


def clean_document_citations(input_file) -> tuple:
    """
    Main orchestration function.
    Reads a document from a path or file buffer (e.g. uploaded file in Streamlit) [1],
    removes citations/references, and returns a tuple:
    (output_bytes_io_buffer, list_of_logs)
    """
    doc = Document(input_file)
    logs = []

    # Step 1: Remove references section
    _, ref_log = remove_references_and_after(doc)
    logs.append(ref_log)

    # Step 2: Clean main body paragraphs
    body_cleaned_count = 0
    for p in doc.paragraphs:
        if clean_paragraph(p):
            body_cleaned_count += 1
    logs.append(f"Processed body: Cleaned citations in {body_cleaned_count} paragraphs.")

    # Step 3: Clean paragraphs inside tables
    table_cleaned_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if clean_paragraph(p):
                        table_cleaned_count += 1
    logs.append(f"Processed tables: Cleaned citations in {table_cleaned_count} table cells.")

    # Save to a memory buffer instead of writing directly to local disk
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)

    return output_buffer, logs
