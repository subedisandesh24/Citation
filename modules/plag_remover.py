import io
import os
import re
import torch
import nltk
from nltk.tokenize import sent_tokenize
from docx import Document
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

# --- PARAPHRASING MODEL LOGIC WRAPPER ---

class ParaphraserEngine:
    """
    Manages the deep learning paraphraser model and tokenizer.
    Loads resources on initialization and handles sentence-level generation.
    """
    def __init__(self, model_name="humarin/chatgpt_paraphraser_on_T5_base"):
        self._init_nltk_resources()
        
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.model_name = model_name
        
        # Load Hugging Face tokenizer and model
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
        self.model = AutoModelForSeq2SeqLM.from_pretrained(self.model_name).to(self.device)

    def _init_nltk_resources(self):
        """Safely ensures the required NLTK resources are loaded in read-only environments."""
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt', quiet=True)
        try:
            nltk.data.find('tokenizers/punkt_tab')
        except LookupError:
            nltk.download('punkt_tab', quiet=True)

    def get_paraphrase(self, text: str) -> str:
        """Runs the sequence-to-sequence model inference on a single sentence."""
        prefix = "paraphrase: " + text
        encoding = self.tokenizer(prefix, padding=True, truncation=True, max_length=512, return_tensors="pt")
        input_ids = encoding["input_ids"].to(self.device)
        attention_masks = encoding["attention_mask"].to(self.device)

        outputs = self.model.generate(
            input_ids=input_ids,
            attention_mask=attention_masks,
            max_length=512,
            do_sample=True,
            top_k=60,
            top_p=0.92,
            temperature=1.15,
            repetition_penalty=1.2,
            early_stopping=True,
            num_return_sequences=1
        )
        return self.tokenizer.decode(outputs[0], skip_special_tokens=True, clean_up_tokenization_spaces=True)


# --- FORMAT CHECKING HELPERS ---

def should_skip(para) -> bool:
    """
    Analyzes formatting characteristics of a paragraph. 
    Returns True if paragraph is a caption, heading, or entirely bolded text and should be bypassed.
    """
    text = para.text.strip()

    if not text:
        return True

    # Check for Table/Figure captions
    if re.match(r'^(Table|Figure|Fig\.)\s+\d+', text, re.IGNORECASE):
        return True

    # Check for Heading Styles
    if para.style.name.startswith('Heading'):
        return True

    # Check for Bold Text
    is_bold = True
    has_text = False
    for run in para.runs:
        if run.text.strip():
            has_text = True
            if not run.bold:
                is_bold = False
                break

    if has_text and is_bold:
        return True

    return False


# --- STREAM-COMPATIBLE ORCHESTRATION PIPELINE ---

def run_paraphrase_document(input_file, engine=None, progress_callback=None) -> tuple:
    """
    Parses a document structure (file path or bytes buffer) sentence-by-sentence,
    performs paraphrasing using the transformers engine, and yields the output buffer.
    
    Args:
        input_file: String path or file-like stream (BytesIO) from Streamlit [1].
        engine: An optional pre-initialized instance of ParaphraserEngine. 
                Providing this prevents reloading the model on every function call.
        progress_callback: Callable function(current_step, total_steps, text_preview).
        
    Returns:
        (output_buffer_bytes, changes_count, log_list)
    """
    # Use pre-initialized engine if provided; otherwise load a new one
    if engine is None:
        engine = ParaphraserEngine()
    
    doc = Document(input_file)
    total_paras = len(doc.paragraphs)
    changes_count = 0
    logs = [f"Found {total_paras} paragraphs. Initiating deep learning pipeline..."]

    for index, para in enumerate(doc.paragraphs):
        original_text = para.text.strip()

        # Update progress tracking if callback provided
        if progress_callback:
            preview = original_text[:40] + "..." if len(original_text) > 40 else original_text
            progress_callback(index + 1, total_paras, preview)

        if should_skip(para):
            continue

        if len(original_text) < 20:
            continue

        try:
            sentences = sent_tokenize(original_text)
            new_sentences = []

            for sent in sentences:
                if len(sent) < 5:
                    new_sentences.append(sent)
                    continue

                new_sent = engine.get_paraphrase(sent)
                new_sentences.append(new_sent)

            final_text = " ".join(new_sentences)

            if final_text != original_text:
                para.text = final_text
                changes_count += 1

        except Exception as e:
            logs.append(f"Error on paragraph index {index}: {str(e)}")

    # Save finalized document structure into an in-memory buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)

    logs.append(f"Processing Complete! {changes_count} paragraphs modified.")
    return output_buffer, changes_count, logs
