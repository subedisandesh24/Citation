import io
import os
import re
import time
import requests
from docx import Document
from docx.shared import Inches, Pt

# --- Strict APA 7th Edition Formatting, Filtering & Verification Engine ---

class StrictAPA50CitationEngine:
    def __init__(self, email="github-citations@example.com"):
        # Crossref recommends providing an email for polite API usage
        self.headers = {'User-Agent': f'StreamlitAPAReferenceGenerator/3.0 (mailto:{email})'}
        self.metadata_cache = {}
        self.cited_dois = set()

    def search_crossref_multi(self, query):
        """Searches Crossref and returns up to 10 candidate papers."""
        clean_query = query.strip()
        if clean_query in self.metadata_cache:
            return self.metadata_cache[clean_query]
            
        url = "https://api.crossref.org/works"
        params = {"query": clean_query, "rows": 10}
        
        try:
            response = requests.get(url, params=params, headers=self.headers, timeout=12)
            if response.status_code == 200:
                items = response.json().get('message', {}).get('items', [])
                self.metadata_cache[clean_query] = items
                return items
        except Exception:
            pass
        return []

    def validate_and_extract_metadata(self, item):
        """
        Validates paper metadata strictly according to APA 7th Edition requirements.
        Returns (is_valid, validation_details)
        """
        if not item:
            return False, {"reason": "Empty payload"}
            
        details = {
            "title_status": "✗ Missing",
            "author_status": "✗ Missing",
            "date_status": "✗ Missing",
            "doi_status": "✗ Missing",
            "source_status": "✗ Missing",
            "type": item.get('type', 'unknown'),
            "title": "N/A",
            "authors": "N/A",
            "date": "N/A",
            "doi": "N/A",
            "journal": "N/A",
            "missing_fields": []
        }
        
        # 1. Validate Title
        title_list = item.get('title', [])
        if title_list and title_list[0].strip():
            details["title_status"] = "✓ Present"
            details["title"] = title_list[0].strip()
        else:
            details["missing_fields"].append("Title")
            
        # 2. Validate Authors
        authors = item.get('author', [])
        valid_authors = []
        if authors:
            for auth in authors:
                if auth.get('family') or auth.get('name') or auth.get('given'):
                    valid_authors.append(auth)
            if valid_authors:
                details["author_status"] = "✓ Present"
                details["authors"] = self.format_authors_apa(valid_authors)
            else:
                details["missing_fields"].append("Author Name")
        else:
            details["missing_fields"].append("Authors List")
                
        # 3. Validate Publication Date (Year)
        year = "n.d."
        for key in ['published-print', 'published-online', 'issued', 'created']:
            date_info = item.get(key)
            if date_info and 'date-parts' in date_info:
                try:
                    extracted_year = date_info['date-parts'][0][0]
                    if extracted_year:
                        year = str(extracted_year)
                        break
                except (IndexError, TypeError):
                    continue
        if year != "n.d.":
            details["date_status"] = "✓ Present"
            details["date"] = year
        else:
            details["missing_fields"].append("Publication Date")
            
        # 4. Validate DOI
        doi = item.get('DOI', '')
        if doi:
            details["doi_status"] = "✓ Present"
            details["doi"] = doi
        else:
            details["missing_fields"].append("DOI")
            
        # 5. Validate Source (Journal or Publisher)
        container_list = item.get('container-title', [])
        publisher = item.get('publisher', '')
        if container_list and container_list[0].strip():
            details["source_status"] = "✓ Present"
            details["journal"] = container_list[0].strip()
        elif publisher and publisher.strip():
            details["source_status"] = "✓ Present"
            details["journal"] = publisher.strip()
        else:
            details["missing_fields"].append("Source (Journal/Publisher Name)")

        # Strict Pass Criteria
        is_valid = len(details["missing_fields"]) == 0
        return is_valid, details

    def select_best_paper(self, candidates):
        """
        Prioritizes journal-articles, evaluates them against criteria, 
        and discards incomplete entries.
        """
        if not candidates:
            return None, None, ["No query results returned from database."]
            
        sorted_candidates = sorted(
            candidates, 
            key=lambda x: 0 if x.get('type') == 'journal-article' else 1
        )
        
        discard_audit_trail = []
        best_paper = None
        best_details = None
        
        for idx, item in enumerate(sorted_candidates):
            is_valid, details = self.validate_and_extract_metadata(item)
            if is_valid:
                best_paper = item
                best_details = details
                break
            else:
                title_preview = item.get('title', ['Untitled'])[0][:45]
                missing_str = ", ".join(details["missing_fields"])
                work_type = item.get('type', 'unknown').upper()
                discard_audit_trail.append(
                    f"Candidate {idx+1} ({work_type}): '{title_preview}...' discarded due to missing: {missing_str}."
                )
                
        return best_paper, best_details, discard_audit_trail

    def to_sentence_case(self, title):
        """Formats article titles into sentence case as required by APA 7th Edition rules."""
        if not title:
            return ""
        title = re.sub('<[^<]+?>', '', title).strip()
        parts = title.split(':')
        formatted_parts = []
        for part in parts:
            part = part.strip()
            if not part:
                continue
            words = part.split()
            if words:
                first_word = words[0][0].upper() + words[0][1:] if len(words[0]) > 1 else words[0].upper()
                rest = []
                for w in words[1:]:
                    if w.isupper() and len(w) > 1:
                        rest.append(w)
                    elif any(c.isupper() for c in w[1:]):
                        rest.append(w)
                    else:
                        rest.append(w.lower())
                formatted_parts.append(first_word + (" " + " ".join(rest) if rest else ""))
        return ": ".join(formatted_parts)

    def format_authors_apa(self, authors_list):
        """Formats the author string according to APA 7th guidelines."""
        formatted = []
        for author in authors_list:
            family = author.get('family', '')
            given = author.get('given', '')
            if not family:
                name = author.get('name', '')
                if name:
                    formatted.append(name)
                continue
            
            initials = "".join([f"{part[0]}." for part in re.split(r'\s+|-', given) if part])
            initials_clean = ". ".join([p.strip() for p in initials.split('.') if p.strip()]) + "." if initials else ""
            formatted.append(f"{family}, {initials_clean}".strip())
            
        num_authors = len(formatted)
        if num_authors == 0:
            return "Unknown Author"
        elif num_authors == 1:
            return formatted[0]
        elif num_authors == 2:
            return f"{formatted[0]}, & {formatted[1]}"
        elif num_authors <= 20:
            return ", ".join(formatted[:-1]) + f", & {formatted[-1]}"
        else:
            first_19 = formatted[:19]
            last = formatted[-1]
            return ", ".join(first_19) + f", … {last}"

    def get_in_text_citation(self, metadata, doi):
        """Generates parenthetical in-text citations incorporating author count."""
        authors = metadata.get('author', [])
        year = "n.d."
        for key in ['published-print', 'published-online', 'issued', 'created']:
            date_info = metadata.get(key)
            if date_info and 'date-parts' in date_info:
                try:
                    year = str(date_info['date-parts'][0][0])
                    break
                except (IndexError, TypeError):
                    continue

        if not authors:
            author_str = metadata.get('publisher', 'Unknown Author')
        elif len(authors) == 1:
            author_str = authors[0].get('family', authors[0].get('name', 'Unknown'))
        elif len(authors) == 2:
            name1 = authors[0].get('family', 'Unknown')
            name2 = authors[1].get('family', 'Unknown')
            author_str = f"{name1} & {name2}"
        else:
            name1 = authors[0].get('family', 'Unknown')
            author_str = f"{name1} et al."

        return f"({author_str}, {year})"

    def format_apa_reference(self, metadata):
        """Formats the bibliography entry."""
        authors_list = metadata.get('author', [])
        authors_str = self.format_authors_apa(authors_list)
        
        date_str = "n.d."
        for key in ['published-print', 'published-online', 'issued', 'created']:
            date_info = metadata.get(key)
            if date_info and 'date-parts' in date_info:
                try:
                    date_str = str(date_info['date-parts'][0][0])
                    break
                except (IndexError, TypeError):
                    continue
        
        title_list = metadata.get('title', [])
        title = title_list[0] if title_list else "Untitled"
        title_sentence = self.to_sentence_case(title)
        
        container_list = metadata.get('container-title', [])
        container = container_list[0] if container_list else ""
        
        volume = metadata.get('volume', '')
        issue = metadata.get('issue', '')
        page = metadata.get('page', '')
        
        doi = metadata.get('DOI', '')
        doi_url = f"https://doi.org/{doi}" if doi else ""
        
        ref = f"{authors_str} ({date_str})."
        
        if container:
            ref += f" {title_sentence}. <i>{container}</i>"
            if volume:
                ref += f", <i>{volume}</i>"
            if issue:
                ref += f"({issue})"
            if page:
                ref += f", {page}."
            else:
                ref += "."
        else:
            ref += f" <i>{title_sentence}</i>."
            publisher = metadata.get('publisher', '')
            if publisher:
                ref += f" {publisher}."
                
        if doi_url:
            ref += f" {doi_url}"
            
        return ref

    def extract_keywords_from_sentence(self, text):
        """Filters keywords and focuses specifically on Nepalese context."""
        stopwords = {
            "the", "a", "an", "and", "or", "but", "if", "then", "else", "when", "at", "by", "from", 
            "for", "with", "about", "against", "between", "into", "through", "during", "before", 
            "after", "above", "below", "to", "of", "in", "on", "is", "are", "was", "were", "be", 
            "been", "being", "have", "has", "had", "doing", "this", "that", "these", "those"
        }
        transition_words = {
            "moreover", "however", "therefore", "although", "furthermore", "thus", "consequently", 
            "additionally", "instead", "meanwhile", "nevertheless", "nonetheless", "besides", 
            "otherwise", "accordingly", "hence", "likewise", "similarly", "finally", "subsequently",
            "recently", "previously", "indeed", "specifically", "especially", "particularly",
            "clearly", "notably", "importantly", "primarily", "secondly", "thirdly", "lastly"
        }
        words = re.findall(r'\b[a-zA-Z]{4,15}\b', text.lower())
        unique_words = []
        for w in words:
            if w not in stopwords and w not in transition_words and w not in unique_words:
                unique_words.append(w)
                
        query_words = unique_words[:4]
        query = " ".join(query_words)
        
        if "nepal" not in query.lower():
            query += " Nepal"
            
        return query


# --- Document Processing Routing adapted for File Buffers ---

def generate_nepal_citations(input_file, target_citations=50, progress_callback=None):
    """
    Reads a document from a file path or in-memory BytesIO stream, inserts citations, 
    and returns a processed BytesIO document buffer along with verification data.
    """
    engine = StrictAPA50CitationEngine()
    
    # Load document from either path or file-like object
    doc = Document(input_file)
    
    paragraphs_sentences = []
    candidates = []
    
    for p_idx, paragraph in enumerate(doc.paragraphs):
        p_text = paragraph.text.strip()
        if not p_text:
            paragraphs_sentences.append([])
            continue
        
        sentences = re.split(r'(?<=[.!?])\s+', p_text)
        paragraphs_sentences.append(sentences)
        
        for s_idx, s_text in enumerate(sentences):
            clean_s = s_text.strip()
            word_count = len(clean_s.split())
            if word_count >= 8 and not clean_s.endswith('?'):
                candidates.append((p_idx, s_idx, clean_s))
                
    total_candidates = len(candidates)
    if total_candidates == 0:
        # Return empty data structures if no citation candidates are present
        return None, [], [], []
        
    num_to_cite = min(target_citations, total_candidates)
    if total_candidates > num_to_cite:
        selected_candidates = [candidates[int(i * (total_candidates - 1) / (num_to_cite - 1))] for i in range(num_to_cite)]
    else:
        selected_candidates = candidates
        
    references_dict = {}
    verification_data = []
    selected_set = {(c[0], c[1]) for c in selected_candidates}
    
    citation_count = 0
    for p_idx, sentences in enumerate(paragraphs_sentences):
        for s_idx, s_text in enumerate(sentences):
            if (p_idx, s_idx) in selected_set:
                citation_count += 1
                keywords = engine.extract_keywords_from_sentence(s_text)
                
                # Report progress if callback is provided
                if progress_callback:
                    progress_callback(citation_count, num_to_cite, keywords)
                
                candidates_list = engine.search_crossref_multi(keywords)
                best_paper, details, discard_logs = engine.select_best_paper(candidates_list)
                
                time.sleep(0.1) # Small rate limiting pause
                
                if best_paper and details:
                    doi = details['doi']
                    citation_str = engine.get_in_text_citation(best_paper, doi)
                    ref_html = engine.format_apa_reference(best_paper)
                    
                    sort_key = engine.format_authors_apa(best_paper.get('author', []))
                    references_dict[doi] = {"sort_key": sort_key, "text": ref_html}
                    
                    m = re.match(r'^(.*?)([.!?]+)$', s_text.strip())
                    if m:
                        base_text = m.group(1)
                        punctuation = m.group(2)
                        paragraphs_sentences[p_idx][s_idx] = f"{base_text} {citation_str}{punctuation}"
                    else:
                        paragraphs_sentences[p_idx][s_idx] = f"{s_text.strip()} {citation_str}"
                        
                    verification_data.append({
                        "id": citation_count,
                        "context": s_text,
                        "keywords": keywords,
                        "title": details["title"],
                        "author": details["authors"],
                        "date": details["date"],
                        "doi": f"https://doi.org/{doi}",
                        "type": details["type"],
                        "journal": details["journal"],
                        "title_status": details["title_status"],
                        "author_status": details["author_status"],
                        "date_status": details["date_status"],
                        "doi_status": details["doi_status"],
                        "source_status": details["source_status"],
                        "discard_logs": discard_logs,
                        "status": "Passed Validation"
                    })
                else:
                    verification_data.append({
                        "id": citation_count,
                        "context": s_text,
                        "keywords": keywords,
                        "title": "N/A",
                        "author": "N/A",
                        "date": "N/A",
                        "doi": "N/A",
                        "type": "N/A",
                        "journal": "N/A",
                        "title_status": "✗ Discarded",
                        "author_status": "✗ Discarded",
                        "date_status": "✗ Discarded",
                        "doi_status": "✗ Discarded",
                        "source_status": "✗ Discarded",
                        "discard_logs": discard_logs,
                        "status": "Discarded (Incomplete APA)"
                    })
                    
    # Rebuild paragraphs
    processed_paragraphs = []
    for p_idx, sentences in enumerate(paragraphs_sentences):
        if sentences:
            rejoined_text = " ".join(sentences)
            doc.paragraphs[p_idx].text = rejoined_text
            processed_paragraphs.append(rejoined_text)
        else:
            processed_paragraphs.append("")
            
    # Build References section
    sorted_reference_texts = []
    if references_dict:
        doc.add_page_break()
        
        h_paragraph = doc.add_paragraph()
        h_paragraph.alignment = 1
        h_format = h_paragraph.paragraph_format
        h_format.space_before = Pt(24)
        h_format.space_after = Pt(12)
        
        h_run = h_paragraph.add_run("References")
        h_run.bold = True
        h_run.font.name = 'Times New Roman'
        h_run.font.size = Pt(12)
        
        sorted_refs = sorted(references_dict.values(), key=lambda x: x['sort_key'].lower())
        sorted_reference_texts = [ref['text'] for ref in sorted_refs]
        
        for ref in sorted_refs:
            r_paragraph = doc.add_paragraph()
            r_format = r_paragraph.paragraph_format
            r_format.left_indent = Inches(0.5)
            r_format.first_line_indent = Inches(-0.5)
            r_format.line_spacing = 2.0
            r_format.space_after = Pt(6)
            
            tokens = re.split(r'(<i>|</i>)', ref['text'])
            italic_toggle = False
            for token in tokens:
                if token == "<i>":
                    italic_toggle = True
                elif token == "</i>":
                    italic_toggle = False
                else:
                    if token:
                        run = r_paragraph.add_run(token)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.italic = italic_toggle
                        
    # Save the updated document directly to an in-memory buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    return output_buffer, processed_paragraphs, sorted_reference_texts, verification_data
